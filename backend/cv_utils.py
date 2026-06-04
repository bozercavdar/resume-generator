from pathlib import Path
from io import BytesIO
import pdfplumber
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages).strip()


def render_html(resume_data: dict, template_name: str) -> str:
    env = Environment(
        loader=FileSystemLoader(BASE_DIR / "templates"),
        autoescape=True,
    )
    template = env.get_template(template_name)
    return template.render(resume=resume_data)


def html_to_pdf(html_content: str) -> bytes:
    return HTML(string=html_content).write_pdf()


def generate_cover_letter_docx(
    text: str,
    name: str,
    email: str | None = None,
) -> bytes:
    doc = Document()

    header = doc.add_paragraph()
    run = header.add_run(name)
    run.bold = True
    run.font.size = Pt(14)

    if email:
        p = doc.add_paragraph()
        p.add_run(email).italic = True

    doc.add_paragraph()

    for paragraph in text.split("\n\n"):
        p = doc.add_paragraph()
        p.add_run(paragraph).font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
