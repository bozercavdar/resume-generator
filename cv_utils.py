from pathlib import Path
import json
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent


def load_resume_data(path: Path) -> dict:
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def render_html(resume_data: dict, template_name: str) -> str:
    env = Environment(
        loader=FileSystemLoader(BASE_DIR / "templates"),
        autoescape=True,
    )
    template = env.get_template(template_name)
    return template.render(resume=resume_data)


def html_to_pdf(html_content: str, output_path: Path) -> None:
    HTML(string=html_content).write_pdf(str(output_path))

def save_cover_letter_docx(
    text: str,
    output_path: Path,
    name: str,
    email: str | None = None
):
    doc = Document()

    # Name header
    header = doc.add_paragraph()
    run = header.add_run(name)
    run.bold = True
    run.font.size = Pt(14)

    if email:
        p = doc.add_paragraph()
        p.add_run(email).italic = True

    doc.add_paragraph()  # spacer

    # Body paragraphs
    for paragraph in text.split("\n\n"):
        p = doc.add_paragraph()
        p.add_run(paragraph).font.size = Pt(11)

    doc.save(output_path)
